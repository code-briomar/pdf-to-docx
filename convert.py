#!/usr/bin/env python3
"""PDF -> Word converter. See README.md for behavior."""
import argparse
import hashlib
import io
import json
import re
import string
import sys
from collections import Counter
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt

# ---------------------------------------------------------------------------
# PDF opening: password detection + corruption repair (pikepdf)
# ---------------------------------------------------------------------------

def open_pdf(path, password=None):
    try:
        doc = fitz.open(path)
    except Exception as e:
        doc = None
        open_err = e
    else:
        open_err = None

    if doc is not None and doc.is_encrypted:
        if not doc.authenticate(password or ""):
            raise SystemExit(
                f"'{path}' is password-protected. Re-run with --password <pw>."
            )
        return doc

    if doc is not None and open_err is None:
        return doc

    # Failed to open at all -> attempt repair via pikepdf, then retry.
    import pikepdf
    try:
        with pikepdf.open(path, password=password or "") as pdf:
            repaired = io.BytesIO()
            pdf.save(repaired)
        doc = fitz.open(stream=repaired.getvalue(), filetype="pdf")
        print(f"[warn] '{path}' was malformed; repaired via pikepdf before conversion.", file=sys.stderr)
        return doc
    except pikepdf.PasswordError:
        raise SystemExit(
            f"'{path}' is password-protected. Re-run with --password <pw>."
        )
    except Exception as e:
        raise SystemExit(f"'{path}' could not be opened or repaired: {e}")


# ---------------------------------------------------------------------------
# Gibberish sanity check for a supposedly-native text layer
# ---------------------------------------------------------------------------

def gibberish_ratio(text):
    if not text or not text.strip():
        return 1.0
    total = len(text)
    junk = sum(
        1 for c in text
        if not (c.isalnum() or c.isspace() or c in string.punctuation)
    )
    return junk / total


# ---------------------------------------------------------------------------
# Reading-order: gap-based column clustering
# ponytail: heuristic gap clustering, not a learned layout model.
# Upgrade to PPStructure's column detection if columns are irregular/overlapping.
# ---------------------------------------------------------------------------

def order_reading_order(blocks, page_width):
    if not blocks:
        return blocks
    by_x = sorted(blocks, key=lambda b: b["bbox"][0])
    gap_threshold = page_width * 0.08
    columns = [[by_x[0]]]
    prev_x0 = by_x[0]["bbox"][0]
    for b in by_x[1:]:
        x0 = b["bbox"][0]
        if x0 - prev_x0 > gap_threshold:
            columns.append([])
        columns[-1].append(b)
        prev_x0 = x0
    ordered = []
    for col in columns:
        ordered.extend(sorted(col, key=lambda b: b["bbox"][1]))
    return ordered


# ---------------------------------------------------------------------------
# Native (text-layer) page extraction
# ---------------------------------------------------------------------------

def _bbox_overlaps(a, b):
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    return not (ax1 < bx0 or bx1 < ax0 or ay1 < by0 or by1 < ay0)


def extract_native_page(page, body_size):
    blocks = []
    tables = []
    for tab in page.find_tables().tables:
        try:
            rows = tab.extract()
        except Exception:
            continue
        tables.append({"bbox": tab.bbox, "rows": rows})

    table_bboxes = [t["bbox"] for t in tables]

    raw = page.get_text("dict")
    for block in raw.get("blocks", []):
        if block.get("type") != 0:  # not text
            continue
        bbox = tuple(block["bbox"])
        if any(_bbox_overlaps(bbox, tb) for tb in table_bboxes):
            continue
        lines_text = []
        max_size = 0.0
        bold = False
        for line in block.get("lines", []):
            line_text = "".join(span["text"] for span in line["spans"])
            if not line_text.strip():
                continue
            lines_text.append(line_text)
            for span in line["spans"]:
                max_size = max(max_size, span["size"])
                if span["flags"] & 2**4:
                    bold = True
        text = "\n".join(lines_text).strip()
        if not text:
            continue
        level = None
        if max_size > body_size * 1.45:
            level = 1
        elif max_size > body_size * 1.2 or (bold and max_size > body_size * 1.05):
            level = 2
        blocks.append({"bbox": bbox, "text": text, "level": level})

    for t in tables:
        blocks.append({"bbox": t["bbox"], "text": None, "table": t["rows"]})

    ordered = order_reading_order(blocks, page.rect.width)
    return ordered


def compute_body_size(doc):
    sizes = Counter()
    for page in doc:
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line["spans"]:
                    if span["text"].strip():
                        sizes[round(span["size"])] += len(span["text"])
    if not sizes:
        return 11.0
    return float(sizes.most_common(1)[0][0])


# ---------------------------------------------------------------------------
# OCR path: rasterize -> deskew -> PaddleOCR (detection + recognition)
# ponytail: plain PaddleOCR, not PPStructure. PPStructure's layout classifier
# was dropping/duplicating regions and its text recognition was consistently
# worse on dense small text (screenshots, UI exports) than the plain det+rec
# pipeline at the same DPI. Loses PPStructure's table-structure detection on
# scanned pages as a result - a scanned table becomes plain OCR'd paragraphs.
# Revisit if scanned tables become a common case.
# ---------------------------------------------------------------------------

_OCR_ENGINES = {}


def _get_ocr_engine(lang):
    if lang not in _OCR_ENGINES:
        try:
            from paddleocr import PaddleOCR
        except ImportError:
            raise SystemExit(
                "This PDF needs OCR but paddleocr/paddlepaddle are not installed.\n"
                "Run: pip install -r requirements.txt"
            )
        _OCR_ENGINES[lang] = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)
    return _OCR_ENGINES[lang]


def deskew(gray):
    import cv2
    import numpy as np
    thresh = cv2.threshold(cv2.bitwise_not(gray), 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
    coords = np.column_stack(np.where(thresh > 0))
    if coords.size == 0:
        return gray
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    # ponytail: real scan skew is a few degrees at most. minAreaRect can lock onto
    # a sparse/lopsided content mass (e.g. a thin margin bar) and report a large
    # angle that isn't skew at all - ignore those rather than rotate a fine page.
    if abs(angle) < 0.1 or abs(angle) > 15:
        return gray
    h, w = gray.shape
    m = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
    return cv2.warpAffine(gray, m, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)


def ocr_page(page, dpi, lang, min_confidence):
    import cv2
    import numpy as np

    pix = page.get_pixmap(dpi=dpi)
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        img = cv2.cvtColor(img, cv2.COLOR_RGBA2GRAY)
    elif pix.n == 3:
        img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    img = deskew(img)
    bgr = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)

    engine = _get_ocr_engine(lang)
    result = engine.ocr(bgr, cls=True)
    lines = result[0] or []

    # Each line is [[[x,y],[x,y],[x,y],[x,y]], (text, confidence)].
    raw = []
    heights = []
    for points, (text, conf) in lines:
        if not text.strip():
            continue
        xs = [p[0] for p in points]
        ys = [p[1] for p in points]
        bbox = (min(xs), min(ys), max(xs), max(ys))
        raw.append({"bbox": bbox, "text": text, "conf": conf})
        heights.append(bbox[3] - bbox[1])

    if not raw:
        return [], 1.0

    # ponytail: line-height-vs-body-height heuristic for heading level, mirroring
    # compute_body_size's approach on the native-text path. No font metadata is
    # available from OCR, so text height is the only signal we have.
    body_height = Counter(round(h) for h in heights).most_common(1)[0][0] or 1

    blocks = []
    scores = []
    for r in raw:
        h = r["bbox"][3] - r["bbox"][1]
        level = 1 if h > body_height * 1.45 else (2 if h > body_height * 1.2 else None)
        blocks.append({"bbox": r["bbox"], "text": r["text"], "level": level})
        scores.append(r["conf"])

    ordered = order_reading_order(blocks, page.rect.width)
    page_confidence = sum(scores) / len(scores) if scores else 1.0
    return ordered, page_confidence


# ---------------------------------------------------------------------------
# Repeating header/footer detection across the whole document
# ---------------------------------------------------------------------------

def strip_repeating_headers_footers(pages_blocks, page_heights):
    if len(pages_blocks) < 3:
        return pages_blocks
    normalized_counts = Counter()
    for blocks, height in zip(pages_blocks, page_heights):
        for b in blocks:
            if b.get("text") is None:
                continue
            y0 = b["bbox"][1]
            if y0 > height * 0.12 and y0 < height * 0.88:
                continue  # not in header/footer zone
            norm = re.sub(r"\d+", "#", b["text"].strip().lower())
            normalized_counts[norm] += 1

    repeat_threshold = max(2, int(len(pages_blocks) * 0.5))
    repeating = {k for k, v in normalized_counts.items() if v >= repeat_threshold}

    cleaned = []
    for blocks, height in zip(pages_blocks, page_heights):
        kept = []
        for b in blocks:
            if b.get("text") is None:
                kept.append(b)
                continue
            y0 = b["bbox"][1]
            norm = re.sub(r"\d+", "#", b["text"].strip().lower())
            in_zone = y0 <= height * 0.12 or y0 >= height * 0.88
            if in_zone and (norm in repeating or re.fullmatch(r"#+|page\s*#(\s*of\s*#)?", norm)):
                continue
            kept.append(b)
        cleaned.append(kept)
    return cleaned


# ---------------------------------------------------------------------------
# docx assembly
# ---------------------------------------------------------------------------

def build_docx(pages_blocks, output_path):
    doc = Document()
    for blocks in pages_blocks:
        for b in blocks:
            if "table" in b:
                rows = b["table"]
                if not rows or not rows[0]:
                    continue
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c in range(ncols):
                        cell_text = row[c] if c < len(row) else ""
                        table.cell(r, c).text = cell_text or ""
            else:
                text = b["text"]
                level = b.get("level")
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                else:
                    doc.add_paragraph(text)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# Checkpointing (resume an interrupted run)
# ---------------------------------------------------------------------------

def checkpoint_path(output_path):
    return Path(str(output_path) + ".checkpoint.json")


def file_hash(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def load_checkpoint(output_path, input_hash):
    cp = checkpoint_path(output_path)
    if not cp.exists():
        return {}
    try:
        data = json.loads(cp.read_text())
    except Exception:
        return {}
    if data.get("input_hash") != input_hash:
        return {}
    return {int(k): v for k, v in data.get("pages", {}).items()}


def save_checkpoint(output_path, input_hash, pages):
    cp = checkpoint_path(output_path)
    cp.write_text(json.dumps({"input_hash": input_hash, "pages": pages}))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser(description="Convert a PDF to a formatted .docx, locally.")
    ap.add_argument("input", help="input PDF path")
    ap.add_argument("-o", "--output", required=True, help="output .docx path")
    ap.add_argument("--force-ocr", action="store_true", help="OCR every page even if a text layer exists")
    ap.add_argument("--lang", default="en", help="OCR language, default en")
    ap.add_argument("--dpi", type=int, default=400, help="rasterization DPI for scanned pages")
    ap.add_argument("--min-confidence", type=float, default=0.6, help="below this, page is flagged for manual review")
    ap.add_argument("--password", default=None, help="password for an encrypted PDF")
    args = ap.parse_args()

    input_hash = file_hash(args.input)
    doc = open_pdf(args.input, args.password)
    body_size = compute_body_size(doc)

    checkpoint = load_checkpoint(args.output, input_hash)
    page_confidences = {}
    page_heights = []

    for i, page in enumerate(doc):
        page_heights.append(page.rect.height)
        if i in checkpoint:
            page_confidences[i] = checkpoint[i].get("_confidence", 1.0)
            continue

        native_text = page.get_text()
        needs_ocr = args.force_ocr or not native_text.strip()
        if not needs_ocr and gibberish_ratio(native_text) > 0.15:
            needs_ocr = True
            print(f"[warn] page {i + 1}: text layer looks garbled, falling back to OCR.", file=sys.stderr)
        # ponytail: word-count-vs-image-coverage heuristic, not a real content classifier.
        # Catches "screenshot exported to PDF" pages that have a tiny stray text layer
        # (e.g. one leftover run) sitting on top of raster content that get_text() can't see.
        if not needs_ocr and len(native_text.split()) < 15 and page.get_image_info():
            needs_ocr = True
            print(f"[warn] page {i + 1}: sparse text layer over image content, falling back to OCR.", file=sys.stderr)

        if needs_ocr:
            blocks, confidence = ocr_page(page, args.dpi, args.lang, args.min_confidence)
        else:
            blocks = extract_native_page(page, body_size)
            confidence = 1.0

        page_confidences[i] = confidence
        checkpoint[i] = {"blocks": blocks, "_confidence": confidence}
        save_checkpoint(args.output, input_hash, checkpoint)
        print(f"[progress] page {i + 1}/{doc.page_count} done", file=sys.stderr)

    pages_blocks = [checkpoint[i]["blocks"] for i in range(doc.page_count)]
    pages_blocks = strip_repeating_headers_footers(pages_blocks, page_heights)
    build_docx(pages_blocks, args.output)

    log_path = Path(str(args.output) + ".log")
    flagged = [i + 1 for i, c in page_confidences.items() if c < args.min_confidence]
    if flagged:
        log_path.write_text(
            "Pages flagged for manual review (confidence below "
            f"{args.min_confidence}):\n" + "\n".join(f"page {p}" for p in sorted(flagged))
        )
        print(f"[warn] {len(flagged)} page(s) flagged for review -> {log_path}", file=sys.stderr)
    elif log_path.exists():
        log_path.unlink()

    checkpoint_path(args.output).unlink(missing_ok=True)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
